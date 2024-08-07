import tkinter as tk
import sqlite3
from tkinter import messagebox
from tkinter import ttk
from docx import Document
import tkinter as tk
from mail import EmailSystemApp
from p_mail import EmailSystemApp2
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import os

DATABASE_PATH = 'D:\\py-temel\\personelmanagemnt\\pymanagement.db'
def quit_button():
    os._exit(0)

def get_db_connection():
    connection = sqlite3.connect(DATABASE_PATH, timeout=10)  # Increased timeout
    return connection

def open_email_system():
    root.withdraw()  
    email_app_window = tk.Toplevel(root)  
    app = EmailSystemApp(email_app_window)  

    def on_closing():
        email_app_window.destroy()  
        root.deiconify()  

    email_app_window.protocol("WM_DELETE_WINDOW", on_closing)  

def open_email_system2():
    root.withdraw()  
    email_app_window = tk.Toplevel(root)  
    app = EmailSystemApp2(email_app_window)  

    def on_closing():
        email_app_window.destroy()  
        root.deiconify()  

    email_app_window.protocol("WM_DELETE_WINDOW", on_closing)  

def mail_main():
    global root
    root = tk.Tk()
    root.title("Mail")
    root.geometry("300x200")

    welcome_label = tk.Label(root, text="Mail sistemini hoş geldin")
    welcome_label.pack(pady=20)

    open_button = tk.Button(root, text="Giriş Yap", command=open_email_system)
    open_button.pack(pady=20)

    root.mainloop()
def mail_main2():
    global root
    root = tk.Tk()
    root.title("Mail")
    root.geometry("300x200")

    welcome_label = tk.Label(root, text="Mail sistemini hoş geldin")
    welcome_label.pack(pady=20)

    open_button = tk.Button(root, text="Giriş Yap", command=open_email_system2)
    open_button.pack(pady=20)

    root.mainloop()

# Admin Paneli
class AdminPanel:
    def __init__(self, user_text):
        self.user_text = user_text
        self.admin_page = tk.Tk()
        self.admin_page.geometry("1920x1080")
        self.admin_page.title("Admin Page")
        self.dep_entry = None
        self.del_entry = None
        self.listbox = None
        self.listbox2 = None
        self.listbox3 = None
        self.setup_ui()

    def logout2(self):
        self.admin_page.destroy()
        login()

    def setup_ui(self):
        connection = get_db_connection()
        cursor = connection.cursor()
        try:
            cursor.execute("SELECT name, surname, level FROM admin WHERE a_username = ?;", (self.user_text,))
            result = cursor.fetchone()
        finally:
            connection.close()

        head_font = ("Arial", 15)
        label_font = ("Arial", 12)

        # Left panel
        label_frame = tk.Frame(self.admin_page, width=20, border=2, relief="groove")
        label_frame.pack(side="left", fill="y", padx=10, pady=10)
        head_label = tk.Label(label_frame, text="Yönetici Bilgileri", font=head_font)
        head_label.pack(padx=10, pady=10)

        if result:
            name, surname, level = result
            label = tk.Label(label_frame, text=f"Adınız: {name} {surname}", font=label_font)
            label.pack(padx=10, pady=10)
            if level == 1:
                label2 = tk.Label(label_frame, text="Mevkiiniz: CEO", font=label_font).pack(padx=10, pady=10)
            elif level == 2:
                label2 = tk.Label(label_frame, text="Mevkiiniz: Müdür", font=label_font).pack(padx=10, pady=10)
            elif level == 3:
                label2 = tk.Label(label_frame, text="Mevkiiniz: Müdür Yardımcısı", font=label_font).pack(padx=10, pady=10)
            elif level == 4:
                label2 = tk.Label(label_frame, text="Mevkiiniz: Şef", font=label_font).pack(padx=10, pady=10)
            elif level == 5:
                label2 = tk.Label(label_frame, text="Mevkiiniz: Şef Yardımcısı", font=label_font).pack(padx=10, pady=10)
            else:
                label2 = tk.Label(label_frame, text="Mevkiiniz: Bulunamadı.", font=label_font)
                label2.pack(padx=10, pady=10)
        else:
            label = tk.Label(label_frame, text="Kullanıcı bilgileri bulunamadı.", font=label_font)
            label.pack(padx=10, pady=10)

        # Left panel #2
        head2_label = tk.Label(label_frame, text="İşlemler", font=head_font)
        head2_label.pack(padx=10, pady=10)
        label = tk.Label(label_frame, text="Departman oluştur", font=label_font)
        label.pack(padx=10, pady=10)
        self.dep_entry = tk.Entry(label_frame)
        self.dep_entry.pack(padx=10, pady=10)
        dep_button = tk.Button(label_frame, text="Oluştur", command=self.departman_create)
        dep_button.pack(padx=10, pady=10)
        label = tk.Label(label_frame, text="Departman Sil (id yazınız)", font=label_font)
        label.pack(padx=10, pady=10)
        self.del_entry = tk.Entry(label_frame)
        self.del_entry.pack(padx=10)
        del_button = tk.Button(label_frame, text="Sil", command=self.departman_delete)
        del_button.pack(padx=10, pady=10)
        dep_frame = tk.Frame(label_frame, border=2, relief="groove")
        dep_frame.pack(fill="x", padx=10, pady=10)
        dep_label = tk.Label(dep_frame, text="Departmanlar", font=head_font)
        dep_label.pack(pady=10)
        self.listbox = tk.Listbox(dep_frame, height=10, width=50)
        self.listbox.pack(side="left", fill="both", expand=True, padx=10, pady=10)
        scrollbar = ttk.Scrollbar(dep_frame, orient=tk.VERTICAL, command=self.listbox.yview)
        scrollbar.pack(side="right", fill="y")
        self.listbox['yscrollcommand'] = scrollbar.set
        self.refresh_listbox()

        # Center panel -> personel panel
        label_frame_r = tk.Frame(self.admin_page, width=20, border=2, relief="groove")
        label_frame_r.pack(side="top", fill="x", padx=10, pady=10)

        # Personeller listbox
        listbox_container2 = tk.Frame(label_frame_r)
        listbox_container2.pack(padx=10, pady=10, fill="both", expand=True)
        label = tk.Label(listbox_container2, text="Personeller", font=head_font)
        label.pack(padx=10)
        self.listbox2 = tk.Listbox(listbox_container2, height=10, width=30)
        self.listbox2.pack(side="left", fill="both", expand=True)
        scrollbar = ttk.Scrollbar(listbox_container2, orient=tk.VERTICAL, command=self.listbox2.yview)
        scrollbar.pack(side="right", fill="y")
        self.listbox2['yscrollcommand'] = scrollbar.set
        self.listbox2.bind('<<ListboxSelect>>', self.on_personel_select)  # Bind selection event

        # PERSONEL GÜNCELLEME 
        self.admin_page = tk.Frame(self.admin_page)
        self.admin_page.pack(padx=10, pady=10, fill="both", expand=True)

        self.canvas = tk.Canvas(self.admin_page)
        self.scrollbar = tk.Scrollbar(self.admin_page, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=self.scrollbar.set)

        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")

        self.personel_details_frame = tk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.personel_details_frame, anchor="nw")

        self.personel_entries = []
        labels = ["ID:", "Adı Soyadı:", "Kullanıcı Adı:", "Yaş:", "Departman:", "Maaş:", "Seviye:", "İş Kalitesi:", "İş Tamamlama Süresi:", "Hedeflere Ulaşma Oranı:", "Takım Çalışması:", "Problem Çözme:"]

        for i, label_text in enumerate(labels):
            row = i // 2  # Her iki etiketi bir satıra yerleştir
            col = (i % 2) * 2  # Etiket ve giriş alanını yan yana yerleştir

            tk.Label(self.personel_details_frame, text=label_text).grid(row=row, column=col, padx=5, pady=5, sticky="e")
            entry = tk.Entry(self.personel_details_frame)
            entry.grid(row=row, column=col + 1, padx=5, pady=5, sticky="we")
            self.personel_entries.append(entry)

        # Update the scroll region of the canvas
        self.personel_details_frame.bind("<Configure>", self.on_frame_configure)
    
        update_button = tk.Button(dep_frame, text="Güncelle", command=self.update_personel)
        update_button.pack(padx=10, pady=10)
        tk.Label(label_frame_r, text="Mail", font=head_font).pack(padx=10, pady=10)
        tk.Button(label_frame_r, text="Mail Oluştur", command=mail_main).pack(padx=10, pady=10)
        tk.Button(dep_frame, text="Çıkış Yap", command=self.logout2).pack(padx=10, pady=10)
        tk.Label(label_frame_r, text="Rapor Üret", font=head_font).pack(padx=10, pady=10)
        button_p_rapor = tk.Button(label_frame_r, text="Personellerin Çıktısını Al", command=self.per_export).pack(padx=10, pady=10)
        button_p_rapor2 = tk.Button(label_frame_r, text="Personellerin Performansının Çıktısını Al", command=self.personel_perfonmans_export).pack(padx=10, pady=10)

        # İK Personelleri listbox
        listbox_container3 = tk.Frame(label_frame_r)
        listbox_container3.pack(padx=10, pady=10, fill="both", expand=True)
        label = tk.Label(listbox_container3, text="İnsan Kaynakları Personelleri ", font=head_font)
        label.pack(padx=10)
        self.listbox3 = tk.Listbox(listbox_container3, height=10, width=30)
        self.listbox3.pack(side="left", fill="both", expand=True)
        scrollbar = ttk.Scrollbar(listbox_container3, orient=tk.VERTICAL, command=self.listbox3.yview)
        scrollbar.pack(side="right", fill="y")
        self.listbox3['yscrollcommand'] = scrollbar.set
        label = tk.Label(label_frame_r, text="İnsan Kaynaklarına Personel Ata", font=head_font).pack(padx=10, pady=10)
        button_hr_add = tk.Button(label_frame_r, text="Ekle", command=HRPanel).pack(padx=10, pady=10)
        self.refresh_listbox3()
        self.refresh_listbox2()

        # Seviyeye göre içerik eklenebilir
        if result:
            if level == 1:
                self.setup_level_1_features()
            elif level == 2:
                self.setup_level_2_features()
            elif level == 3:
                self.setup_level_3_features()
            elif level == 4:
                self.setup_level_4_features()
            elif level == 5:
                self.setup_level_5_features()
            else:
                label2 = tk.Label(label_frame_r, text="Mevkiiniz: Bulunamadı.", font=label_font).pack(padx=10, pady=10)
        else:
            label = tk.Label(label_frame, text="Kullanıcı bilgileri bulunamadı.", font=label_font)
            label.pack(padx=10, pady=10)

    def on_frame_configure(self, event):
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def on_personel_select(self, event):
        selected_index = self.listbox2.curselection()
        if selected_index:
            selected_id = self.listbox2.get(selected_index).split(",")[0].split(":")[1].strip()
            connection = get_db_connection()
            cursor = connection.cursor()
            cursor.execute("SELECT id, name_surname, username_p, age, department, salary, level, work_quality, job_completion_time, goals_rate, teamwork, problem_solving FROM personnel WHERE id = ?", (selected_id,))
            personel = cursor.fetchone()
            connection.close()
            if personel:
                for i, entry in enumerate(self.personel_entries):
                    entry.delete(0, tk.END)
                    entry.insert(0, personel[i])

    def update_personel(self):
        selected_index = self.listbox2.curselection()
        if not selected_index:
            messagebox.showwarning("Uyarı", "Lütfen güncellenecek personeli seçin.")
            return

        selected_id = self.listbox2.get(selected_index).split(",")[0].split(":")[1].strip()
        updated_values = [entry.get() for entry in self.personel_entries]

        connection = get_db_connection()
        cursor = connection.cursor()
        cursor.execute("""
            UPDATE personnel
            SET id = ?, name_surname = ?, username_p = ?, age = ?, department = ?, salary = ?, level = ?, work_quality = ?, job_completion_time = ?, goals_rate = ?, teamwork = ?, problem_solving = ?
            WHERE id = ?
        """, (*updated_values, selected_id))
        connection.commit()
        connection.close()
        self.refresh_listbox2()
        messagebox.showinfo("Bilgi", "Personel bilgileri başarıyla güncellendi.")

    def departman_create(self):
        departman = self.dep_entry.get()
        if not departman:
            messagebox.showwarning("Uyarı", "Lütfen departman adını girin.")
            return
        connection = get_db_connection()
        cursor = connection.cursor()
        cursor.execute("INSERT INTO departman (name) VALUES (?)", (departman,))
        connection.commit()
        connection.close()
        self.refresh_listbox()

    def departman_delete(self):
        departman_id = self.del_entry.get()
        if not departman_id:
            messagebox.showwarning("Uyarı", "Lütfen departman ID'sini girin.")
            return
        connection = get_db_connection()
        cursor = connection.cursor()
        cursor.execute("DELETE FROM departman WHERE id = ?", (departman_id,))
        connection.commit()
        connection.close()
        self.refresh_listbox()

    def per_export(self):
        document = Document()
        document.add_heading('Personel Raporu', 0)
        connection = get_db_connection()
        cursor = connection.cursor()
        cursor.execute("SELECT * FROM personnel")
        personel_list = cursor.fetchall()
        connection.close()
        for personel in personel_list:
            document.add_paragraph(f"ID: {personel[0]}, Ad: {personel[1]}, Soyad: {personel[2]}, E-posta: {personel[3]}")
        document.save('personel_raporu.docx')
        messagebox.showinfo("Bilgi", "Rapor başarıyla oluşturuldu.")

    def personel_perfonmans_export(self):
        document = Document()
        document.add_heading('Personel Performans Raporu', 0)
        connection = get_db_connection()
        cursor = connection.cursor()
        cursor.execute("SELECT * FROM personnel")
        personel_list = cursor.fetchall()
        connection.close()
        for personel in personel_list:
            # Örnek performans verisi eklenebilir
            document.add_paragraph(f"ID: {personel[0]}, Ad: {personel[1]}, Soyad: {personel[2]}, Performans: TBD")
        document.save('personel_performans_raporu.docx')
        messagebox.showinfo("Bilgi", "Performans raporu başarıyla oluşturuldu.")

    def refresh_listbox(self):
        self.listbox.delete(0, tk.END)
        connection = get_db_connection()
        cursor = connection.cursor()
        try:
            cursor.execute("SELECT id, d_name FROM departman")
            result = cursor.fetchall()
            for item in result:
                self.listbox.insert(tk.END, f"id: {item[0]} {item[1]}")
        finally:
            connection.close()

    def refresh_listbox2(self):
        self.listbox2.delete(0, tk.END)
        connection = get_db_connection()
        cursor = connection.cursor()
        try:
            cursor.execute("SELECT id, name_surname, username_p, age, department, salary, level, work_quality, job_completion_time, goals_rate, teamwork, problem_solving FROM personnel")
            result = cursor.fetchall()
            for item in result:
                self.listbox2.insert(tk.END, f"Id:{item[0]}, İsim: {item[1]}, Kullanıcı Adı: {item[2]}, Yaş: {item[3]}, Departman: {item[4]}, Maaş: {item[5]}, Seviye: {item[6]}, İş Kalitesi: {item[7]}, İş Tamamlama Süresi: {item[8]}, Hedeflere Ulaşma Oranı: {item[9]}, Takım Çalışması: {item[10]}, Problem Çözme: {item[11]}")
        finally:
            connection.close()

    def refresh_listbox3(self):
        self.listbox3.delete(0, tk.END)
        connection = get_db_connection()
        cursor = connection.cursor()
        try:
            cursor.execute("SELECT id, hr_username,hr_password,hr_name,hr_surname,hr_level,hr_salary FROM hr")
            result = cursor.fetchall()
            for item in result:
                self.listbox3.insert(tk.END, f"id: {item[0]} {item[1]} {item[2]} {item[3]} {item[4]} {item[5]} {item[6]}")
        finally:
            connection.close()

    def setup_level_1_features(self):
        pass  # CEO'ya özel işlemler burada

    def setup_level_2_features(self):
        pass  # Müdür'e özel işlemler burada

    def setup_level_3_features(self):
        pass  # Müdür Yardımcısı'na özel işlemler burada

    def setup_level_4_features(self):
        pass  # Şef'e özel işlemler burada

    def setup_level_5_features(self):
        pass  # Şef Yardımcısı'na özel işlemler burada

def admin_login():
    pass_text = entry_pass.get()
    user_text = entry_user.get()
    connection = get_db_connection()
    cursor = connection.cursor()
    try:
        cursor.execute("SELECT a_username, a_password FROM admin WHERE a_username = ? AND a_password = ?;", (user_text, pass_text))
        if cursor.fetchone():
            messagebox.showinfo("Uyarı!", "Giriş Başarılı.")
            login_page.destroy()
            AdminPanel(user_text)
        else:
            messagebox.showerror("Giriş Hatası!", "Kullanıcı adı veya şifre yanlış")
    finally:
        connection.close()

def hr_login():
    pass_text = entry_pass.get()
    user_text = entry_user.get()
    connection = get_db_connection()
    cursor = connection.cursor()
    try:
        cursor.execute("SELECT hr_username, hr_password FROM hr WHERE hr_username = ? AND hr_password = ?;", (user_text, pass_text))
        if cursor.fetchone():
            messagebox.showinfo("Uyarı!", "Giriş Başarılı.")
            login_page.destroy()
            hr_main_panel(user_text)
        else:
            messagebox.showerror("Giriş Hatası!", "Kullanıcı adı veya şifre yanlış")
    finally:
        connection.close()

class hr_main_panel:
    def __init__(self, user_text):
        self.user_text = user_text
        self.hr_page2 = tk.Tk()
        self.hr_page2.geometry("1920x1080")
        self.hr_page2.title("İnsan Kaynakları Sayfası")
        self.hr_name_entry = None
        self.hr_username_entry = None
        self.hr_password_entry = None
        self.del_entry = None
        self.hr_level_entry = None
        self.hr_salary_entry = None
        self.pr_age_entry = None
        self.hr_departmant_entry = None
        self.listbox2 = None
        self.update_entries = {}
        self.performance_entries = {}

        self.setup_ui()

    def setup_ui(self):
        global listbox2, listbox
        head_font = ("Arial", 15)
        label_font = ("Arial", 12)

        label_frame = tk.Frame(self.hr_page2, width=20, border=2, relief="groove")
        label_frame.pack(side="left", fill="y", padx=10, pady=10)

        label = tk.Label(label_frame, text="Personel Kayıt Formu", font=head_font)
        label.pack(pady=10)

        tk.Label(label_frame, text="Adı Soyadı:", font=label_font).pack(pady=5)
        self.hr_name_entry = tk.Entry(label_frame)
        self.hr_name_entry.pack(pady=5)

        tk.Label(label_frame, text="Yaşı:", font=label_font).pack(pady=5)
        self.pr_age_entry = tk.Entry(label_frame)
        self.pr_age_entry.pack(pady=5)

        tk.Label(label_frame, text="Kullanıcı Adı:", font=label_font).pack(pady=5)
        self.hr_username_entry = tk.Entry(label_frame)
        self.hr_username_entry.pack(pady=5)

        tk.Label(label_frame, text="Şifre:", font=label_font).pack(pady=5)
        self.hr_password_entry = tk.Entry(label_frame, show="*")
        self.hr_password_entry.pack(pady=5)

        tk.Label(label_frame, text="Seviye:", font=label_font).pack(pady=5)
        self.hr_level_entry = tk.Entry(label_frame)
        self.hr_level_entry.pack(pady=5)

        tk.Label(label_frame, text="Departman:", font=label_font).pack(pady=5)
        self.hr_departmant_entry = tk.Entry(label_frame)
        self.hr_departmant_entry.pack(pady=5)

        tk.Label(label_frame, text="Maaş:", font=label_font).pack(pady=5)
        self.hr_salary_entry = tk.Entry(label_frame)
        self.hr_salary_entry.pack(pady=5)

        tk.Button(label_frame, text="Ekle", command=self.per_add).pack(pady=20)

        # Department list below personnel registration form
        dep_frame = tk.Frame(label_frame, border=2, relief="groove")
        dep_frame.pack(fill="x", padx=10, pady=10)

        dep_label = tk.Label(dep_frame, text="Departmanlar", font=head_font)
        dep_label.pack(pady=10)

        self.listbox = tk.Listbox(dep_frame, height=10, width=50)
        self.listbox.pack(side="left", fill="both", expand=True, padx=10, pady=10)

        scrollbar = ttk.Scrollbar(dep_frame, orient=tk.VERTICAL, command=self.listbox.yview)
        scrollbar.pack(side="right", fill="y")
        self.listbox['yscrollcommand'] = scrollbar.set
        tk.Button(dep_frame, text="Mail Oluştur", command=mail_main).pack(padx=10, pady=10)
        tk.Button(dep_frame, text="Çıkış Yap", command=self.logout2).pack(padx=10, pady=10)
        self.refresh_listbox()

        # Center panel for personnel operations
        label_frame_r = tk.Frame(self.hr_page2, width=20, border=2, relief="groove")
        label_frame_r.pack(side="top", fill="x", padx=10, pady=10)

        listbox_container2 = tk.Frame(label_frame_r)
        listbox_container2.pack(padx=10, pady=10, fill="both", expand=True)
        label = tk.Label(listbox_container2, text="Personeller & Personel İşlemleri", font=head_font)
        label.pack(padx=10)
        self.listbox2 = tk.Listbox(listbox_container2, height=10, width=30)
        self.listbox2.pack(side="left", fill="both", expand=True)
        scrollbar = ttk.Scrollbar(listbox_container2, orient=tk.VERTICAL, command=self.listbox2.yview)
        scrollbar.pack(side="right", fill="y")
        self.listbox2['yscrollcommand'] = scrollbar.set

        # Bind the selection event to the on_personnel_select function
        self.listbox2.bind('<<ListboxSelect>>', self.on_personnel_select)

        left_frame = tk.Frame(label_frame_r, relief="groove")
        left_frame.pack(side="left", fill="y", padx=10, pady=10)

        label = tk.Label(left_frame, text="Personel Çıkart", font=head_font)
        label.pack(padx=10, pady=10)

        self.del_entry = tk.Entry(left_frame)
        self.del_entry.pack(padx=10, pady=10)

        del_button = tk.Button(left_frame, text="Sil", command=self.per_delete)
        del_button.pack(padx=10, pady=10)

        # Right frame
        right_frame = tk.Frame(label_frame_r, relief="groove")
        right_frame.pack(side="top", fill="x", padx=10, pady=10)

        update_label = tk.Label(right_frame, text="Personel Güncelle | Yukarıdan Personel Seçmeyi Unutma !", font=head_font)
        update_label.pack(padx=10, pady=10)
        field_frame = tk.Frame(right_frame)
        field_frame.pack(pady=5, fill="x", padx=10)

        fields = ["ID", "Kullanıcı Adı", "Yaş", "Departman", "Maaş", "Seviye"]
        self.update_entries = {}
        for field in fields:
            field_label = tk.Label(field_frame, text=field, font=label_font)
            field_label.pack(side="left", padx=5)
            entry = tk.Entry(field_frame)
            entry.pack(side="left", fill="x", expand=True, padx=5)
            self.update_entries[field] = entry

        update_button = tk.Button(right_frame, text="Güncelle", command=self.per_update)
        update_button.pack(padx=10, pady=10)

        # Performance frame
        performance_frame = tk.Frame(self.hr_page2, relief="groove")
        performance_frame.pack(side="top", fill="x", padx=10, pady=10)

        performance_label = tk.Label(performance_frame, text="Personel Performans Güncelleme", font=head_font)
        performance_label.pack(padx=10, pady=10)
        performance_fields = ["İş Kalitesi", "İş Tamamlama Süresi", "Hedeflere Ulaşma Oranı", "Takım Çalışması", "Problem Çözme"]
        self.performance_entries = {}
        for field in performance_fields:
            field_label = tk.Label(performance_frame, text=field, font=label_font)
            field_label.pack(side="left", padx=5)
            entry = tk.Entry(performance_frame)
            entry.pack(side="left", fill="x", expand=True, padx=5)
            self.performance_entries[field] = entry

        performance_update_button = tk.Button(performance_frame, text="Performans Güncelle", command=self.performance_update)
        performance_update_button.pack(padx=10, pady=10)

        label_frame_b = tk.Frame(self.hr_page2, width=20, border=2, relief="groove")
        label_frame_b.pack(side="top", fill="x", padx=10, pady=10)
        label2 = tk.Label(label_frame_b, text="Rapor Üret", font=head_font).pack(padx=10, pady=10)
        button_p_rapor = tk.Button(label_frame_b, text="Personellerin Çıktısını Al", command=self.per_export).pack(padx=10, pady=10)
        button_p_rapor2 = tk.Button(label_frame_b, text="Personellerin Performansının Çıktısını Al", command=self.personel_perfonmans_export).pack(padx=10, pady=10)

        self.refresh_listbox2()
        self.hr_page2.mainloop()

    def refresh_listbox(self):
        self.listbox.delete(0, tk.END)
        connection = get_db_connection()
        cursor = connection.cursor()
        try:
            cursor.execute("SELECT id, d_name FROM departman")
            result = cursor.fetchall()
            for item in result:
                self.listbox.insert(tk.END, f"id: {item[0]} {item[1]}")
        finally:
            connection.close()

    def refresh_listbox2(self):
        self.listbox2.delete(0, tk.END)
        connection = get_db_connection()
        cursor = connection.cursor()
        try:
            cursor.execute("SELECT id, name_surname, username_p, age, department, salary, level, work_quality, job_completion_time, goals_rate, teamwork, problem_solving FROM personnel")
            result = cursor.fetchall()
            for item in result:
                self.listbox2.insert(tk.END, f"Id:{item[0]}, İsim: {item[1]}, Kullanıcı Adı: {item[2]}, Yaş: {item[3]}, Departman: {item[4]}, Maaş: {item[5]}, Seviye: {item[6]}, İş Kalitesi: {item[7]}, İş Tamamlama Süresi: {item[8]}, Hedeflere Ulaşma Oranı: {item[9]}, Takım Çalışması: {item[10]}, Problem Çözme: {item[11]}")
        finally:
            connection.close()

    def logout2(self):
        self.hr_page2.destroy()
        login()

    def on_personnel_select(self, event):
        selected_index = self.listbox2.curselection()
        if selected_index:
            selected_personnel = self.listbox2.get(selected_index)
            personnel_data = selected_personnel.split(", ")
            personnel_info = {}
            for data in personnel_data:
                key, value = data.split(":")
                personnel_info[key.strip()] = value.strip()

            self.update_entries["ID"].config(state='normal')
            self.update_entries["ID"].delete(0, tk.END)
            self.update_entries["ID"].insert(0, personnel_info["Id"])
            self.update_entries["ID"].config(state='readonly')

            self.update_entries["Kullanıcı Adı"].delete(0, tk.END)
            self.update_entries["Kullanıcı Adı"].insert(0, personnel_info["Kullanıcı Adı"])

            self.update_entries["Yaş"].delete(0, tk.END)
            self.update_entries["Yaş"].insert(0, personnel_info["Yaş"])

            self.update_entries["Departman"].delete(0, tk.END)
            self.update_entries["Departman"].insert(0, personnel_info["Departman"])

            self.update_entries["Maaş"].delete(0, tk.END)
            self.update_entries["Maaş"].insert(0, personnel_info["Maaş"])

            self.update_entries["Seviye"].delete(0, tk.END)
            self.update_entries["Seviye"].insert(0, personnel_info["Seviye"])

            self.performance_entries["İş Kalitesi"].delete(0, tk.END)
            self.performance_entries["İş Kalitesi"].insert(0, personnel_info["İş Kalitesi"])

            self.performance_entries["İş Tamamlama Süresi"].delete(0, tk.END)
            self.performance_entries["İş Tamamlama Süresi"].insert(0, personnel_info["İş Tamamlama Süresi"])

            self.performance_entries["Hedeflere Ulaşma Oranı"].delete(0, tk.END)
            self.performance_entries["Hedeflere Ulaşma Oranı"].insert(0, personnel_info["Hedeflere Ulaşma Oranı"])

            self.performance_entries["Takım Çalışması"].delete(0, tk.END)
            self.performance_entries["Takım Çalışması"].insert(0, personnel_info["Takım Çalışması"])

            self.performance_entries["Problem Çözme"].delete(0, tk.END)
            self.performance_entries["Problem Çözme"].insert(0, personnel_info["Problem Çözme"])

    def per_add(self):
        pr_name = self.hr_name_entry.get()
        pr_age = self.pr_age_entry.get()
        pr_username = self.hr_username_entry.get()
        pr_password = self.hr_password_entry.get()
        pr_level = self.hr_level_entry.get()
        pr_department = self.hr_departmant_entry.get()
        pr_salary = self.hr_salary_entry.get()

        if not all([pr_name, pr_age, pr_username, pr_password, pr_level, pr_department, pr_salary]):
            messagebox.showerror("Hata!", "Lütfen tüm alanları doldurunuz.")
            return

        connection = get_db_connection()
        cursor = connection.cursor()
        try:
            cursor.execute("""
                INSERT INTO personnel (name_surname, username_p, password_p, age, department, salary, level)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (pr_name, pr_username, pr_password, pr_age, pr_department, pr_salary, pr_level))
            connection.commit()
            self.refresh_listbox2()
            messagebox.showinfo("Başarılı", "Personel başarıyla eklendi.")
        except sqlite3.Error as e:
            messagebox.showerror("Hata!", f"Veritabanı hatası: {e}")
        finally:
            connection.close()

    def per_delete(self):
        p_id = self.del_entry.get()
        if not p_id:
            messagebox.showerror("Hata!", "Lütfen geçerli bir personel Id giriniz.")
            return

        connection = get_db_connection()
        cursor = connection.cursor()
        try:
            cursor.execute("DELETE FROM personnel WHERE id = ?", (p_id,))
            if cursor.rowcount == 0:
                messagebox.showerror("Hata!", "Belirtilen Id'ye sahip personel bulunamadı.")
            else:
                connection.commit()
                self.refresh_listbox2()
                messagebox.showinfo("Başarılı", "Personel başarıyla silindi.")
        except sqlite3.Error as e:
            messagebox.showerror("Hata!", f"Veritabanı hatası: {e}")
        finally:
            connection.close()

    def per_update(self):
        personnel_id = self.update_entries["ID"].get()
        if not personnel_id:
            messagebox.showerror("Hata!", "Güncellenecek personel seçiniz.")
            return

        updates = {field: self.update_entries[field].get() for field in self.update_entries if field != "ID"}

        if not any(updates.values()):
            messagebox.showerror("Hata!", "Lütfen en az bir alanı doldurun.")
            return

        connection = get_db_connection()
        cursor = connection.cursor()
        try:
            for field, value in updates.items():
                if value:
                    column_name = {
                        "Kullanıcı Adı": "username_p",
                        "Yaş": "age",
                        "Departman": "department",
                        "Maaş": "salary",
                        "Seviye": "level"
                    }[field]
                    cursor.execute(f"UPDATE personnel SET {column_name} = ? WHERE id = ?", (value, personnel_id))
            connection.commit()
            self.refresh_listbox2()
            messagebox.showinfo("Başarılı", "Personel bilgileri başarıyla güncellendi.")
        except sqlite3.Error as e:
            messagebox.showerror("Hata!", f"Veritabanı hatası: {e}")
        finally:
            connection.close()


    def performance_update(self):
        personnel_id = self.update_entries["ID"].get()
        if not personnel_id:
            messagebox.showerror("Hata!", "Güncellenecek personel seçiniz.")
            return

        updates = {field: self.performance_entries[field].get() for field in self.performance_entries}

        if not any(updates.values()):
            messagebox.showerror("Hata!", "Lütfen en az bir alanı doldurun.")
            return

        connection = get_db_connection()
        cursor = connection.cursor()
        try:
            for field, value in updates.items():
                if value:
                    column_name = {
                        "İş Kalitesi": "work_quality",
                        "İş Tamamlama Süresi": "job_completion_time",
                        "Hedeflere Ulaşma Oranı": "goals_rate",
                        "Takım Çalışması": "teamwork",
                        "Problem Çözme": "problem_solving"
                    }[field]
                    cursor.execute(f"UPDATE personnel SET {column_name} = ? WHERE id = ?", (value, personnel_id))
            connection.commit()
            self.refresh_listbox2()
            messagebox.showinfo("Başarılı", "Personel performans bilgileri başarıyla güncellendi.")
        except sqlite3.Error as e:
            messagebox.showerror("Hata!", f"Veritabanı hatası: {e}")
        finally:
            connection.close()
    def per_export(self):
        doc = Document()
        items = self.listbox2.get(0, self.listbox2.size() - 1)
        for item in items:
            doc.add_paragraph(item)

        doc.save('D:\\py-temel\\personelmanagemnt\\personel_list.docx')
        messagebox.showinfo("Başarılı", "Personel listesi dışa aktarıldı.")

    def perfomance_calculate(self):
        connection = get_db_connection()
        cursor = connection.cursor()
        cursor.execute("SELECT id, name_surname, username_p, age, department, salary, level, work_quality, job_completion_time, goals_rate, teamwork, problem_solving FROM personnel")
        scores = cursor.fetchall()
        criteria_weights = {
            'work_quality': 0.3,
            'job_completion_time': 0.3,
            'goals_rate': 0.2,
            'teamwork': 0.1,
            'problem_solving': 0.1
        }
        
        total_scores = {}
        if scores:
            for row in scores:
                personnel_id = row[0]
                total_score = 0  
                for i, criterion in enumerate(criteria_weights.keys(), start=7):
                    try:
                        score = float(row[i])  
                        weight = criteria_weights[criterion] 
                        total_score += score * weight
                    except ValueError as e:
                        print(f"Error converting score to float: {e}")
                total_scores[personnel_id] = total_score
        return total_scores

    def personel_perfonmans_export(self):
        doc = Document()
        items = self.listbox2.get(0, self.listbox2.size() - 1)
        total_scores = self.perfomance_calculate()
        doc.add_paragraph("1-5: Kötü | 5-7: İyi | 7-10: Çok iyi")

        for item in items:
            personnel_id = int(item.split(",")[0].split(":")[1].strip())
            total_score = total_scores.get(personnel_id, "N/A")
            paragraph_text = f"{item}, Total Score: {total_score}"
            doc.add_paragraph(paragraph_text)
            if total_score < 5:
                doc.add_paragraph("Perfonmansı kötü")
            elif 5 < total_score < 7:
                doc.add_paragraph("Perfonmansı iyi")
            elif 7 < total_score < 10:
                doc.add_paragraph("Perfonmansı çok iyi")
            else :
                doc.add_paragraph(" Hata")

        try:
            doc.save('D:\\py-temel\\personelmanagemnt\\personel_perfonmans_list.docx')
            messagebox.showinfo("Başarılı", "Personel perfonmans listesi dışa aktarıldı.")
        except PermissionError as e:
            messagebox.showerror("Hata!", f"Dosya kaydedilemedi: {e}")
        except Exception as e:
            messagebox.showerror("Hata!", f"Beklenmeyen bir hata oluştu: {e}")

class HRPanel:
    def __init__(self):
        self.hr_page = tk.Tk()
        self.hr_page.geometry("1920x1080")
        self.hr_page.title("İnsan Kaynakları Personeli Ekleme Sayfası")
        
        self.hr_name_entry = None
        self.hr_username_entry = None
        self.hr_password_entry = None
        self.hr_level_entry = None
        self.hr_salary_entry = None
        
        self.setup_ui()
        
    def setup_ui(self):
        head_font = ("Arial", 15)
        label_font = ("Arial", 12)
        
        # Entry fields and labels
        label = tk.Label(self.hr_page, text="İnsan Kaynakları Kayıt Formu", font=head_font)
        label.pack(pady=10)

        tk.Label(self.hr_page, text="Adı Soyadı:", font=label_font).pack(pady=5)
        self.hr_name_entry = tk.Entry(self.hr_page)
        self.hr_name_entry.pack(pady=5)

        tk.Label(self.hr_page, text="Kullanıcı Adı:", font=label_font).pack(pady=5)
        self.hr_username_entry = tk.Entry(self.hr_page)
        self.hr_username_entry.pack(pady=5)

        tk.Label(self.hr_page, text="Şifre:", font=label_font).pack(pady=5)
        self.hr_password_entry = tk.Entry(self.hr_page, show="*")
        self.hr_password_entry.pack(pady=5)

        tk.Label(self.hr_page, text="Seviye:", font=label_font).pack(pady=5)
        self.hr_level_entry = tk.Entry(self.hr_page)
        self.hr_level_entry.pack(pady=5)

        tk.Label(self.hr_page, text="Maaş:", font=label_font).pack(pady=5)
        self.hr_salary_entry = tk.Entry(self.hr_page)
        self.hr_salary_entry.pack(pady=5)

        tk.Button(self.hr_page, text="Ekle", command=self.hr_add).pack(pady=20)

        self.hr_page.mainloop()

    def hr_add(self):
        hr_name = self.hr_name_entry.get()
        hr_username = self.hr_username_entry.get()
        hr_password = self.hr_password_entry.get()
        hr_level = self.hr_level_entry.get()
        hr_salary = self.hr_salary_entry.get()

        if not hr_name or not hr_username or not hr_password or not hr_level or not hr_salary:
            messagebox.showerror("Hata!", "Tüm alanları doldurunuz.")
            return

        connection = get_db_connection()
        cursor = connection.cursor()
        try:
            cursor.execute("""
                INSERT INTO hr (hr_name, hr_username, hr_password, hr_level, hr_salary)
                VALUES (?, ?, ?, ?, ?)
            """, (hr_name, hr_username, hr_password, hr_level, hr_salary))
            connection.commit()
            messagebox.showinfo("Başarılı", "İnsan Kaynakları personeli başarıyla eklendi.")
        except sqlite3.Error as e:
            messagebox.showerror("Hata!", f"Bir hata oluştu: {e}")
        finally:
            connection.close()

def home_login():
    conn = get_db_connection()
    cursor = conn.cursor()

    username = entry_user.get()
    password = entry_pass.get()

    cursor.execute("SELECT * FROM personnel WHERE username_p=? AND password_p=?", (username, password))
    user = cursor.fetchone()
    
    if user:
        id, name_surname, username_p, password_p, age, department, salary, level, work_quality, job_completion_time, goals_rate, participation, teamwork, problem_solving = user
        login_page.destroy()
        show_home_page(name_surname, username_p, age, department, salary, level, work_quality, job_completion_time, goals_rate, participation, teamwork, problem_solving)
    else:
        messagebox.showerror("Hata", "Kullanıcı adı veya şifre yanlış.")

    conn.close()

def show_home_page(name_surname, username_p, age, department, salary, level, work_quality, job_completion_time, goals_rate, participation, teamwork, problem_solving):
    global home_page
    home_page = tk.Tk()
    home_page.geometry("1920x1080")
    home_page.title("Personel Sayfası")

    # Ana Frame
    main_frame = tk.Frame(home_page)
    main_frame.pack(fill=tk.BOTH, expand=True)

    # Sol Frame: Bilgiler
    left_frame = tk.Frame(main_frame, width=600, bg='lightgrey')
    left_frame.pack(side=tk.LEFT, fill=tk.Y)

    label = tk.Label(left_frame, text="Personel Bilgileri", font=("Arial", 20), bg='lightgrey')
    label.pack(pady=20)

    info_text = (
        f"Ad ve Soyad: {name_surname}\n"
        f"Kullanıcı Adı: {username_p}\n"
        f"Yaş: {age}\n"
        f"Departman: {department}\n"
        f"Maaş: {salary}\n"
        f"Seviye: {level}\n"
        f"Çalışma Kalitesi: {work_quality}\n"
        f"İş Tamamlanma Süresi: {job_completion_time}\n"
        f"Hedef Oranı: {goals_rate}\n"
        f"Katılım: {participation}\n"
        f"Takım Çalışması: {teamwork}\n"
        f"Problemi Çözme: {problem_solving}"
    )
    
    info_label = tk.Label(left_frame, text=info_text, font=("Arial", 16), bg='lightgrey', justify=tk.LEFT)
    info_label.pack(padx=20, pady=10, anchor='w')
    tk.Button(left_frame, text="Mail Oluştur", command=mail_main2).pack(padx=10, pady=10)
    tk.Button(left_frame, text="Çıkış Yap", command=logout).pack(padx=10, pady=10)

    # Orta Üst Frame: Grafik
    graph_frame = tk.Frame(main_frame)
    graph_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

    # Grafik oluşturma
    fig = Figure(figsize=(10, 6), dpi=100)
    ax = fig.add_subplot(111)

    # Performans verilerini çizgi grafiği olarak gösterme
    criteria = ['Çalışma Kalitesi', 'İş Tamamlanma Süresi', 'Hedef Oranı', 'Takım Çalışması', 'Problemi Çözme']
    performance_values = [float(work_quality), float(job_completion_time), float(goals_rate), float(teamwork), float(problem_solving)]

    ax.plot(criteria, performance_values, marker='o', linestyle='-', color='b', label='Performans Değeri')
    ax.set_xlabel('Kriterler')
    ax.set_ylabel('Değer')
    ax.set_title('Personel Performans Çizgi Grafiği')
    ax.legend()

    # Grafiği Tkinter penceresine yerleştirme
    canvas = FigureCanvasTkAgg(fig, master=graph_frame)
    canvas.draw()
    canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    home_page.mainloop()

def logout():
    global home_page
    home_page.destroy() 
    login()  
def login():
    global entry_user, entry_pass, login_page
    login_page = tk.Tk()
    login_page.geometry("400x300+{}+{}".format(int(login_page.winfo_screenwidth()/2 - 150), int(login_page.winfo_screenheight()/2 - 100)))
    login_page.title("Giriş Sayfası")
    label = tk.Label(login_page, text="Kullanıcı Adı:")
    label.pack()
    entry_user = tk.Entry(login_page)
    entry_user.pack()
    label = tk.Label(login_page, text="Şifre:")
    label.pack()
    entry_pass = tk.Entry(login_page, show="*")
    entry_pass.pack()
    button_login = tk.Button(login_page, text="Giriş Yap", command=home_login)
    button_login.pack(pady=10)
    button_register = tk.Button(login_page, text="İnsan Kaynakları Girişi", command=hr_login)
    button_register.pack(pady=10)
    button_admin = tk.Button(login_page, text="Admin Girişi", command=admin_login)
    button_admin.pack(pady=10)
    login_page.mainloop()


login()